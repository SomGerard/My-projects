import docx, re
from docx.shared import RGBColor
#store the number of times script has been executed


#getIndex: sets an end index for a slice when an iterable is passed to it
# with the start index and marker for the end
def getIndex(startIndex, textList, endMarker):
    pos=startIndex
    while textList[pos] != endMarker:
        pos+=1
    return pos


#seperate strings: gets all the lines in a text file into a list of strings
#seperate the list into individual lists using the end marker and return the list of lists
def seperateStrings(filename):
    inputFile = open(filename)
    #inputFile = open('student letters_3.txt')                                 #open the text file to be read from
    fullText= inputFile.readlines()                                           #get all lines into a list of strings
    inputFile.close()
    filteredText = [line for line in fullText if line != ' \n' and line !='\n'] #use a list comprehension to selectively add elements to the list removing all blank lines
    subList = []
    stringList=[]
    startSlice=0                                   #mark the beginning of the 1st slice of the list; filteredText
    endSlice = filteredText.index('end\n')         #mark the end of the 1st slice of the list; filteredText
    iterationLimit = filteredText.count('end\n')   #set the end of the for loop
    for i in range (iterationLimit):
        subList=filteredText[startSlice : endSlice]    #set the value of the sublist element to the current slice
        stringList.append(subList)                      #add the current sublist element to the end of the main list
        startSlice = endSlice+1                          #set the new start location for the next slice
        if startSlice == len(filteredText):
            break
        endSlice = getIndex (startSlice, filteredText, 'end\n')           #get the index at which to end each slice
    return stringList

#addColouredRun: adds run text to a paragraph using a specific colour     
def addColouredRun (paraObj, colour, runText):
    begin = runText.find(colour)
    end = begin+len(colour)
    if colour == 'Red':
        run = paraObj.add_run(runText[begin:end])
        font =run.font
        font.color.rgb = RGBColor(0xff, 0x00, 0x00)     #make the word Red appear in red colour
        run.bold=True
    elif colour == 'Amber':
        run = paraObj.add_run(runText[begin:end])
        font = run.font
        font.color.rgb = RGBColor(0xff, 0xbf, 0x00)   # make the word Amber appear in amber colour
        run.bold = True
    elif colour == 'Green':
        run = paraObj.add_run(runText[begin:end])
        font = run.font
        font.color.rgb = RGBColor(0x00, 0x80, 0x00)   #make the word green appear in Green colour 
        run.bold=True
    else:
        print('Invalid colour entered!')
    return end



#writeDocx: write the result letters of the students
def writeDocx (list_of_strings, result_file_name):
    colourRegex = re.compile(r'Red|Amber|Green')
    doc = docx.Document('result_letter.docx')
    for item in list_of_strings:
        para1 = doc.add_paragraph(item[0][:len(item[0])-1])
        para1.style = 'result letter style'
        para1.runs[0].bold = True
        para2 = doc.add_paragraph(item[1][:len(item[1])-1])
        para2.style = 'result letter style'
        #paragraph 3 line 1
        para3 = doc.add_paragraph(item[2][:len(item[2])-1])
        para3.runs[0].bold = True
        para3.runs[0].add_break()
        #paragraph 3 line 2
        endIndex = addColouredRun(para3, 'Red', item[3])
        run_2=para3.add_run(item[3][endIndex:len(item[3])-1])
        run_2.add_break()
        #pargraph 3 line 3
        endIndex = addColouredRun(para3, 'Amber', item[4])
        run_4 =para3.add_run(item[4][endIndex:len(item[4])-1])
        run_4.add_break()
        #paragraph 3 line 4
        endIndex = addColouredRun(para3, 'Green', item[5])
        run_6=para3.add_run(item[5][endIndex:len(item[5])-1])
        para3.style = 'result letter style'
        #paragraph 4 line 1
        para4 = doc.add_paragraph(item[6][:len(item[6])-1])
        para4.runs[0].bold = True
        para4.runs[0].add_break()
        #add subjects and grades
        runCount = 0
        for i in range (7,14):
            col = colourRegex.search(item[i])
            if col == None:
                break
            color = col.group()
            begin = item[i].find(color)
            end = begin+len(color)
            para4.add_run(item[i][:begin])
            addColouredRun(para4, color, item[i][begin:end])
            runCount +=2
            para4.runs[runCount].add_break()
        para4.style = 'result letter style'    
        begin = i
        end = item.index('Kind Regards,\n')
        #paragraph 5
        para5 = doc.add_paragraph(item[begin])
        if len(item[begin:end]) ==2 :
            para5.runs[0].add_break()
            para5.add_run(item[begin+1])
        para5.style = 'result letter style'
        #paragraph 6
        para6 = doc.add_paragraph(item[end])
        para6.runs[0].add_break()
        para6.add_run(item[end+1])
        para6.runs[1].add_break()
        para6.add_run(item[end+2])
        para6.style = 'result letter style'
        para6.runs[2].add_break(docx.enum.text.WD_BREAK.PAGE)

    doc.save(result_file_name)
        
    
if __name__ == '__main__':
    print('Getting input from text file...........\n\n')
    inputText = seperateStrings('student letters_3.txt')
    print('Writing students letters...............\n\n')
    writeDocx(inputText)
    print('Completed!')

scriptRunTimes=0
    
        


scriptRunTimes=0
